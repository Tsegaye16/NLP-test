import fitz  # PyMuPDF
import pdfplumber
from pathlib import Path
import io
import logging
from typing import Dict, List, Optional, Tuple, Any, Union
import hashlib
import unicodedata
from datetime import datetime
from dataclasses import dataclass
import json

from src.utils.logger import logger
from config.settings import PATHS

@dataclass
class DocumentMetadata:
    """Metadata for a historical document"""
    document_id: str
    filename: str
    filepath: Path
    file_size: int
    file_hash: str
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    page_count: int = 0
    word_count: int = 0
    character_count: int = 0
    language: str = "en"
    period: Optional[str] = None
    source_type: Optional[str] = None
    title: Optional[str] = None
    author: Optional[str] = None
    year: Optional[int] = None
    keywords: List[str] = None
    
    def __post_init__(self):
        if self.keywords is None:
            self.keywords = []
    
    def to_dict(self) -> Dict:
        """Convert metadata to dictionary"""
        return {
            "document_id": self.document_id,
            "filename": self.filename,
            "file_size": self.file_size,
            "file_hash": self.file_hash,
            "created_date": self.created_date.isoformat() if self.created_date else None,
            "modified_date": self.modified_date.isoformat() if self.modified_date else None,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "character_count": self.character_count,
            "language": self.language,
            "period": self.period,
            "source_type": self.source_type,
            "title": self.title,
            "author": self.author,
            "year": self.year,
            "keywords": self.keywords
        }

class DocumentLoader:
    """Load and process various document formats with comprehensive error handling"""
    
    def __init__(self, data_dir: Optional[Path] = None):
        self.data_dir = data_dir or PATHS.raw_documents
        self.supported_formats = ['.pdf', '.txt', '.docx', '.html', '.htm']
        self.documents = {}
        self.metadata = {}
        
        logger.log_step("DocumentLoader initialized", details={
            "data_dir": str(self.data_dir),
            "supported_formats": self.supported_formats
        })
    
    def discover_documents(self) -> List[Path]:
        """Discover all supported documents in the data directory"""
        documents = []
        
        for format_ext in self.supported_formats:
            found = list(self.data_dir.glob(f"*{format_ext}"))
            documents.extend(found)
            
            logger.log_step("Document discovery", details={
                "format": format_ext,
                "count": len(found)
            })
        
        logger.log_step("Total documents discovered", details={
            "total": len(documents)
        })
        
        return documents
    
    def load_document(self, filepath: Path) -> Tuple[Optional[str], Optional[DocumentMetadata]]:
        """Load a single document with comprehensive error handling"""
        document_id = self._generate_document_id(filepath)
        
        with logger.time_step("load_document", document_id):
            try:
                # Check if file exists
                if not filepath.exists():
                    logger.log_error("FileNotFound", f"File not found: {filepath}", document_id)
                    return None, None
                
                # Read file based on format
                if filepath.suffix.lower() == '.pdf':
                    text, metadata = self._load_pdf(filepath)
                elif filepath.suffix.lower() == '.txt':
                    text, metadata = self._load_text(filepath)
                elif filepath.suffix.lower() in ['.docx']:
                    text, metadata = self._load_docx(filepath)
                elif filepath.suffix.lower() in ['.html', '.htm']:
                    text, metadata = self._load_html(filepath)
                else:
                    logger.log_error("UnsupportedFormat", 
                                   f"Unsupported format: {filepath.suffix}", document_id)
                    return None, None
                
                if text is None:
                    logger.log_error("LoadFailed", f"Failed to load document: {filepath}", document_id)
                    return None, None
                
                # Create metadata
                doc_metadata = self._extract_metadata(filepath, text, metadata)
                doc_metadata.document_id = document_id
                
                # Store document
                self.documents[document_id] = text
                self.metadata[document_id] = doc_metadata
                
                logger.log_step("Document loaded successfully", document_id, {
                    "word_count": doc_metadata.word_count,
                    "page_count": doc_metadata.page_count,
                    "period": doc_metadata.period
                })
                
                return text, doc_metadata
                
            except Exception as e:
                logger.log_error("LoadException", f"Error loading document: {str(e)}", 
                               document_id, e)
                return None, None
    
    def _load_pdf(self, filepath: Path) -> Tuple[Optional[str], Dict]:
        """Load PDF document using multiple methods for robustness"""
        metadata = {}
        text_parts = []
        
        try:
            # Method 1: Try PyMuPDF (fitz) first
            with fitz.open(filepath) as doc:
                metadata['page_count'] = len(doc)
                # Try to get PDF version, but it may not be available
                try:
                    metadata['pdf_version'] = getattr(doc, 'pdf_version', None)
                except:
                    metadata['pdf_version'] = None
                
                for page_num, page in enumerate(doc, start=1):
                    page_text = page.get_text()
                    
                    if page_text.strip():  # Only add if there's content
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}\n")
                    
                    # Extract metadata from first page if available
                    if page_num == 1:
                        try:
                            meta = doc.metadata
                            if meta:
                                metadata.update(meta)
                        except:
                            pass  # Metadata extraction is optional
            
            if text_parts:
                return "\n".join(text_parts), metadata
            
            # Method 2: If PyMuPDF fails, try pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}\n")
                
                if text_parts:
                    return "\n".join(text_parts), metadata
            
            logger.log_step("PDF extraction warning", details={
                "file": filepath.name,
                "message": "No text extracted from PDF, might be scanned/image-based"
            })
            
            return None, metadata
            
        except Exception as e:
            logger.log_error("PDFLoadError", f"Error loading PDF: {str(e)}", 
                           exception=e, details={"file": filepath.name})
            return None, metadata
    
    def _load_text(self, filepath: Path) -> Tuple[Optional[str], Dict]:
        """Load plain text document"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    # Validate text
                    if self._is_valid_text(text):
                        return text, {}
                    
                except UnicodeDecodeError:
                    continue
            
            logger.log_error("TextEncodingError", 
                           f"Could not decode text file with any encoding: {filepath}")
            return None, {}
            
        except Exception as e:
            logger.log_error("TextLoadError", f"Error loading text file: {str(e)}", 
                           exception=e, details={"file": filepath.name})
            return None, {}
    
    def _load_docx(self, filepath: Path) -> Tuple[Optional[str], Dict]:
        """Load DOCX document"""
        try:
            import docx
            
            doc = docx.Document(filepath)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            return text if text.strip() else None, {}
            
        except ImportError:
            logger.log_error("ImportError", "python-docx not installed", 
                           details={"file": filepath.name})
            return None, {}
        except Exception as e:
            logger.log_error("DocxLoadError", f"Error loading DOCX: {str(e)}", 
                           exception=e, details={"file": filepath.name})
            return None, {}
    
    def _load_html(self, filepath: Path) -> Tuple[Optional[str], Dict]:
        """Load HTML document"""
        try:
            from bs4 import BeautifulSoup
            
            with open(filepath, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'lxml')
            
            # Remove script and style elements
            for script in soup(["script", "style", "header", "footer", "nav"]):
                script.decompose()
            
            # Get text
            text = soup.get_text(separator='\n', strip=True)
            
            # Clean up excessive newlines
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            text = '\n'.join(lines)
            
            return text if text.strip() else None, {}
            
        except ImportError:
            logger.log_error("ImportError", "BeautifulSoup/lxml not installed", 
                           details={"file": filepath.name})
            return None, {}
        except Exception as e:
            logger.log_error("HtmlLoadError", f"Error loading HTML: {str(e)}", 
                           exception=e, details={"file": filepath.name})
            return None, {}
    
    def _extract_metadata(self, filepath: Path, text: str, 
                         extra_metadata: Dict) -> DocumentMetadata:
        """Extract comprehensive metadata from document"""
        # Generate file hash
        file_hash = self._calculate_file_hash(filepath)
        
        # Get file stats
        stat = filepath.stat()
        created_date = datetime.fromtimestamp(stat.st_ctime)
        modified_date = datetime.fromtimestamp(stat.st_mtime)
        
        # Basic text statistics
        words = text.split()
        word_count = len(words)
        character_count = len(text)
        
        # Try to extract year from filename or text
        year = self._extract_year(filepath, text)
        
        # Infer period based on year
        period = self._infer_period(year)
        
        # Try to extract title and author
        title = self._extract_title(filepath, text)
        author = self._extract_author(filepath, text, extra_metadata)
        
        # Infer source type
        source_type = self._infer_source_type(text)
        
        # Extract keywords (basic)
        keywords = self._extract_keywords(text)
        
        metadata = DocumentMetadata(
            document_id="",  # Will be set by caller
            filename=filepath.name,
            filepath=filepath,
            file_size=stat.st_size,
            file_hash=file_hash,
            created_date=created_date,
            modified_date=modified_date,
            page_count=extra_metadata.get('page_count', 0),
            word_count=word_count,
            character_count=character_count,
            year=year,
            period=period,
            source_type=source_type,
            title=title,
            author=author,
            keywords=keywords
        )
        
        return metadata
    
    def _calculate_file_hash(self, filepath: Path) -> str:
        """Calculate MD5 hash of file"""
        hash_md5 = hashlib.md5()
        with open(filepath, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _extract_year(self, filepath: Path, text: str) -> Optional[int]:
        """Extract year from filename and text"""
        import re
        
        # Try filename first
        filename = filepath.stem
        year_match = re.search(r'\b(1[0-9]{3}|20[0-9]{2})\b', filename)
        if year_match:
            return int(year_match.group())
        
        # Try text (first 2000 characters)
        sample = text[:2000]
        year_match = re.search(r'\b(1[0-9]{3}|20[0-9]{2})\b', sample)
        if year_match:
            return int(year_match.group())
        
        return None
    
    def _infer_period(self, year: Optional[int]) -> Optional[str]:
        """Infer historical period based on year"""
        if year is None:
            return None
        
        if year < 1500:
            return "Medieval"
        elif 1500 <= year < 1700:
            return "Early Modern"
        elif 1700 <= year < 1800:
            return "18th Century"
        elif 1800 <= year < 1900:
            return "19th Century"
        elif 1900 <= year < 1918:
            return "Early 20th Century"
        elif 1918 <= year < 1945:
            return "Interwar Period"
        elif 1945 <= year < 1990:
            return "Cold War"
        elif 1990 <= year < 2000:
            return "Late 20th Century"
        else:
            return "21st Century"
    
    def _extract_title(self, filepath: Path, text: str) -> Optional[str]:
        """Extract title from filename or text"""
        # Clean filename
        title = filepath.stem.replace('_', ' ').replace('-', ' ').title()
        
        # If text starts with something that looks like a title
        first_lines = text.split('\n')[:5]
        for line in first_lines:
            line = line.strip()
            if (len(line) < 100 and 
                line.istitle() and 
                len(line.split()) <= 10 and
                not line.endswith(('.', '!', '?'))):
                return line
        
        return title
    
    def _extract_author(self, filepath: Path, text: str, 
                       extra_metadata: Dict) -> Optional[str]:
        """Extract author information"""
        # Check PDF metadata
        if 'author' in extra_metadata and extra_metadata['author']:
            return extra_metadata['author']
        
        # Look for common author patterns in text
        import re
        patterns = [
            r'by\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
            r'Author:\s*([^\n]+)',
            r'Written by\s*([^\n]+)'
        ]
        
        first_page = text[:2000]
        for pattern in patterns:
            match = re.search(pattern, first_page, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _infer_source_type(self, text: str) -> str:
        """Infer type of historical document"""
        text_lower = text.lower()
        
        # Check for common patterns
        if any(term in text_lower for term in ['constitution', 'amendment', 'article']):
            return "Legal Document"
        elif any(term in text_lower for term in ['speech', 'address', 'oration']):
            return "Speech"
        elif any(term in text_lower for term in ['treaty', 'agreement', 'accord']):
            return "Treaty"
        elif any(term in text_lower for term in ['manifesto', 'declaration', 'proclamation']):
            return "Political Document"
        elif any(term in text_lower for term in ['report', 'analysis', 'study']):
            return "Report"
        elif any(term in text_lower for term in ['letter', 'correspondence', 'dispatch']):
            return "Correspondence"
        else:
            return "Historical Text"
    
    def _extract_keywords(self, text: str, top_n: int = 10) -> List[str]:
        """Extract basic keywords from text"""
        from collections import Counter
        import re
        
        # Remove common words and get frequency
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        stop_words = {'that', 'with', 'have', 'this', 'from', 'they', 'were', 'their'}
        
        filtered_words = [w for w in words if w not in stop_words]
        word_freq = Counter(filtered_words)
        
        return [word for word, _ in word_freq.most_common(top_n)]
    
    def _generate_document_id(self, filepath: Path) -> str:
        """Generate unique document ID"""
        # Use hash of filepath + timestamp
        import time
        base_str = f"{filepath.name}_{filepath.stat().st_size}_{time.time()}"
        return hashlib.md5(base_str.encode()).hexdigest()[:12]
    
    def _is_valid_text(self, text: str) -> bool:
        """Validate that text contains meaningful content"""
        if not text or len(text.strip()) < 100:
            return False
        
        # Check if text contains mostly printable characters
        printable_ratio = sum(1 for c in text if c.isprintable() or c.isspace()) / len(text)
        return printable_ratio > 0.9
    
    def load_all_documents(self) -> Dict[str, Tuple[str, DocumentMetadata]]:
        """Load all discovered documents"""
        documents = self.discover_documents()
        loaded_docs = {}
        
        logger.log_step("Starting bulk document loading", details={
            "total_documents": len(documents)
        })
        
        for i, doc_path in enumerate(documents, 1):
            logger.log_step(f"Loading document {i}/{len(documents)}", 
                          details={"file": doc_path.name})
            
            text, metadata = self.load_document(doc_path)
            
            if text and metadata:
                loaded_docs[metadata.document_id] = (text, metadata)
        
        logger.log_step("Bulk loading complete", details={
            "loaded": len(loaded_docs),
            "failed": len(documents) - len(loaded_docs)
        })
        
        return loaded_docs
    
    def save_metadata(self, output_dir: Optional[Path] = None):
        """Save all metadata to JSON file"""
        if output_dir is None:
            output_dir = PATHS.processed_dir
        
        metadata_file = output_dir / "document_metadata.json"
        metadata_list = [meta.to_dict() for meta in self.metadata.values()]
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata_list, f, indent=2, ensure_ascii=False)
        
        logger.log_step("Metadata saved", details={
            "file": str(metadata_file),
            "count": len(metadata_list)
        })